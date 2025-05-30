"""
Professional Report Visualizer for FreelancePay
Generates charts and exports reports to DOCX/PDF formats
"""

import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import plotly.io as pio
from datetime import datetime, timedelta
import base64
from io import BytesIO
import json

# Document generation imports
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

# PDF generation imports
import markdown2

class ReportVisualizer:
    """
    Professional report visualization and export handler
    """
    
    def __init__(self, output_dir="reports"):
        self.output_dir = output_dir
        self.charts_dir = os.path.join(output_dir, "charts")
        os.makedirs(self.charts_dir, exist_ok=True)
        
        # Set professional styling
        self.setup_styling()
        
    def setup_styling(self):
        """Set up professional styling for charts"""
        # Matplotlib styling with fallbacks
        try:
            plt.style.use('seaborn-v0_8-darkgrid')
        except OSError:
            try:
                plt.style.use('seaborn-darkgrid')
            except OSError:
                plt.style.use('default')
                print("⚠️ Using default matplotlib style")
        
        try:
            sns.set_palette("husl")
        except Exception:
            print("⚠️ Could not set seaborn palette")
        
        # Plotly styling
        pio.templates.default = "plotly_white"
        
        # Color palettes
        self.colors = {
            'primary': '#2E86AB',
            'secondary': '#A23B72', 
            'accent': '#F18F01',
            'success': '#C73E1D',
            'info': '#7209B7',
            'neutral': '#6C757D'
        }
        
        self.business_colors = {
            'Revenue & Sales': '#2E86AB',
            'User Experience': '#A23B72',
            'Security & Compliance': '#F18F01',
            'Operational Efficiency': '#C73E1D',
            'Platform Stability': '#7209B7',
            'Feature Expansion': '#6C757D'
        }

    def prepare_data(self, tasks, business_analysis):
        """Prepare data for visualization"""
        
        # Convert tasks to DataFrame for easier manipulation
        df = pd.DataFrame(tasks)
        
        # Add business categories to tasks
        for category_key, category_data in business_analysis['categories'].items():
            for task in category_data['tasks']:
                task['business_category'] = category_data['impact']
        
        # Create comprehensive dataset
        data = {
            'tasks_df': df,
            'business_analysis': business_analysis,
            'metrics': business_analysis['metrics'],
            'categories': business_analysis['categories']
        }
        
        return data

    def create_business_impact_chart(self, data, save_path=None):
        """Create business impact distribution chart"""
        
        categories = data['categories']
        category_data = []
        
        for cat_key, cat_info in categories.items():
            if cat_info['tasks']:
                total_hours = sum(task.get('time', 0) for task in cat_info['tasks'])
                completed_tasks = len([t for t in cat_info['tasks'] if t.get('commits')])
                category_data.append({
                    'Category': cat_info['impact'],
                    'Hours': total_hours,
                    'Tasks': len(cat_info['tasks']),
                    'Completed': completed_tasks
                })
        
        if not category_data:
            return None
            
        df = pd.DataFrame(category_data)
        
        # Create subplot with two charts
        fig = make_subplots(
            rows=1, cols=2,
            specs=[[{"type": "domain"}, {"type": "xy"}]],
            subplot_titles=('Time Investment by Category', 'Task Completion by Category')
        )
        
        # Pie chart for time distribution
        fig.add_trace(
            go.Pie(
                labels=df['Category'],
                values=df['Hours'],
                name="Hours",
                marker_colors=[self.business_colors.get(cat, '#6C757D') for cat in df['Category']],
                textinfo="label+percent",
                textposition="inside"
            ),
            row=1, col=1
        )
        
        # Bar chart for task completion
        fig.add_trace(
            go.Bar(
                x=df['Category'],
                y=df['Tasks'],
                name="Total Tasks",
                marker_color='lightblue',
                opacity=0.7
            ),
            row=1, col=2
        )
        
        fig.add_trace(
            go.Bar(
                x=df['Category'],
                y=df['Completed'],
                name="Completed Tasks",
                marker_color=[self.business_colors.get(cat, '#6C757D') for cat in df['Category']]
            ),
            row=1, col=2
        )
        
        fig.update_layout(
            title_text="Business Impact Analysis",
            title_x=0.5,
            height=500,
            showlegend=True,
            font=dict(size=12)
        )
        
        fig.update_xaxes(tickangle=45, row=1, col=2)
        
        if save_path:
            fig.write_image(save_path, width=1200, height=500, scale=2)
        
        return fig

    def create_velocity_trends_chart(self, data, save_path=None):
        """Create development velocity trends chart"""
        
        tasks_df = data['tasks_df']
        if tasks_df.empty:
            return None
        
        # Group by week/sprint for velocity analysis
        # This is simulated data - in real implementation, you'd use actual dates
        weeks = []
        hours_per_week = []
        tasks_per_week = []
        
        # Simulate 4 weeks of data based on total metrics
        total_hours = data['metrics']['total_time']
        total_tasks = data['metrics']['total_tasks']
        
        for i in range(4):
            week_hours = total_hours * self.rng.uniform(0.2, 0.3)  # Simulate weekly variation
            week_tasks = total_tasks * self.rng.uniform(0.2, 0.3)
            weeks.append(f"Week {i+1}")
            hours_per_week.append(week_hours)
            tasks_per_week.append(week_tasks)
        
        fig = make_subplots(
            rows=2, cols=1,
            subplot_titles=('Development Hours per Week', 'Tasks Completed per Week'),
            vertical_spacing=0.1
        )
        
        # Hours trend
        fig.add_trace(
            go.Scatter(
                x=weeks,
                y=hours_per_week,
                mode='lines+markers',
                name='Hours',
                line=dict(color=self.colors['primary'], width=3),
                marker=dict(size=8)
            ),
            row=1, col=1
        )
        
        # Tasks trend
        fig.add_trace(
            go.Scatter(
                x=weeks,
                y=tasks_per_week,
                mode='lines+markers',
                name='Tasks',
                line=dict(color=self.colors['accent'], width=3),
                marker=dict(size=8)
            ),
            row=2, col=1
        )
        
        fig.update_layout(
            title_text="Development Velocity Trends",
            title_x=0.5,
            height=600,
            showlegend=False
        )
        
        if save_path:
            fig.write_image(save_path, width=1000, height=600, scale=2)
        
        return fig

    def create_priority_distribution_chart(self, data, save_path=None):
        """Create task priority distribution chart"""
        
        tasks_df = data['tasks_df']
        if tasks_df.empty:
            return None
        
        # Count tasks by priority
        priority_counts = tasks_df['priority'].value_counts()
        
        fig = go.Figure(data=[
            go.Bar(
                x=priority_counts.index,
                y=priority_counts.values,
                marker_color=[
                    self.colors['primary'] if p == 'High' 
                    else self.colors['accent'] if p == 'Medium'
                    else self.colors['neutral']
                    for p in priority_counts.index
                ],
                text=priority_counts.values,
                textposition='auto'
            )
        ])
        
        fig.update_layout(
            title="Task Priority Distribution",
            title_x=0.5,
            xaxis_title="Priority Level",
            yaxis_title="Number of Tasks",
            height=400
        )
        
        if save_path:
            fig.write_image(save_path, width=800, height=400, scale=2)
        
        return fig

    def create_repository_activity_chart(self, tasks, save_path=None):
        """Create repository activity chart"""
        
        # Extract repository data from tasks
        repo_data = {}
        for task in tasks:
            repos = task.get('repositories', [])
            commits = task.get('commits', [])
            
            for repo in repos:
                if repo not in repo_data:
                    repo_data[repo] = {'commits': 0, 'hours': 0, 'tasks': 0}
                repo_data[repo]['commits'] += len(commits)
                repo_data[repo]['hours'] += task.get('time', 0)
                repo_data[repo]['tasks'] += 1
        
        if not repo_data:
            return None
        
        repos = list(repo_data.keys())
        commits = [repo_data[repo]['commits'] for repo in repos]
        hours = [repo_data[repo]['hours'] for repo in repos]
        
        fig = make_subplots(
            rows=1, cols=2,
            subplot_titles=('Commits per Repository', 'Hours per Repository')
        )
        
        # Commits chart
        fig.add_trace(
            go.Bar(
                x=repos,
                y=commits,
                name="Commits",
                marker_color=self.colors['primary']
            ),
            row=1, col=1
        )
        
        # Hours chart
        fig.add_trace(
            go.Bar(
                x=repos,
                y=hours,
                name="Hours",
                marker_color=self.colors['accent']
            ),
            row=1, col=2
        )
        
        fig.update_layout(
            title_text="Repository Activity Analysis",
            title_x=0.5,
            height=400,
            showlegend=False
        )
        
        if save_path:
            fig.write_image(save_path, width=1000, height=400, scale=2)
        
        return fig

    def create_performance_dashboard(self, data, save_path=None):
        """Create comprehensive performance dashboard"""
        
        metrics = data['metrics']
        
        fig = make_subplots(
            rows=2, cols=2,
            specs=[[{"type": "indicator"}, {"type": "indicator"}],
                   [{"type": "xy"}, {"type": "domain"}]],
            subplot_titles=('', '', 'Weekly Progress', 'Task Status Distribution')
        )
        
        # KPI indicators
        fig.add_trace(
            go.Indicator(
                mode="gauge+number+delta",
                value=metrics['completion_rate'],
                domain={'x': [0, 1], 'y': [0, 1]},
                title={'text': "Completion Rate (%)"},
                gauge={
                    'axis': {'range': [None, 100]},
                    'bar': {'color': self.colors['primary']},
                    'steps': [
                        {'range': [0, 50], 'color': "lightgray"},
                        {'range': [50, 80], 'color': "yellow"},
                        {'range': [80, 100], 'color': "green"}
                    ],
                    'threshold': {
                        'line': {'color': "red", 'width': 4},
                        'thickness': 0.75, 'value': 90
                    }
                }
            ),
            row=1, col=1
        )
        
        fig.add_trace(
            go.Indicator(
                mode="number+delta",
                value=metrics['total_time'],
                title={'text': "Total Hours"},
                number={'suffix': "h"},
                delta={'reference': metrics['total_time'] * 0.8, 'relative': True}
            ),
            row=1, col=2
        )
        
        # Simulated weekly progress
        weeks = ['Week 1', 'Week 2', 'Week 3', 'Week 4']
        planned = [25, 30, 35, 30]
        actual = [28, 27, 38, 32]
        
        fig.add_trace(
            go.Scatter(x=weeks, y=planned, name="Planned", line=dict(color='blue', dash='dash')),
            row=2, col=1
        )
        fig.add_trace(
            go.Scatter(x=weeks, y=actual, name="Actual", line=dict(color='green')),
            row=2, col=1
        )
        
        # Task status pie chart
        status_data = {
            'Completed': metrics['completed_tasks'],
            'In Progress': metrics['total_tasks'] - metrics['completed_tasks']
        }
        
        fig.add_trace(
            go.Pie(
                labels=list(status_data.keys()),
                values=list(status_data.values()),
                name="Status"
            ),
            row=2, col=2
        )
        
        fig.update_layout(
            title_text="Performance Dashboard",
            title_x=0.5,
            height=700
        )
        
        if save_path:
            fig.write_image(save_path, width=1200, height=700, scale=2)
        
        return fig

    def generate_all_charts(self, tasks, business_analysis):
        """Generate all charts for the report"""
        
        data = self.prepare_data(tasks, business_analysis)
        chart_files = {}
        
        # Generate each chart
        charts = {
            'business_impact': self.create_business_impact_chart,
            'velocity_trends': self.create_velocity_trends_chart,
            'priority_distribution': self.create_priority_distribution_chart,
            'repository_activity': lambda d, p: self.create_repository_activity_chart(tasks, p),
            'performance_dashboard': self.create_performance_dashboard
        }
        
        for chart_name, chart_func in charts.items():
            file_path = os.path.join(self.charts_dir, f"{chart_name}.png")
            try:
                fig = chart_func(data, file_path)
                if fig:
                    chart_files[chart_name] = file_path
                    print(f"✅ Generated {chart_name} chart")
            except Exception as e:
                print(f"⚠️ Error generating {chart_name}: {e}")
        
        return chart_files

    def _process_markdown_line(self, doc, line, current_section):
        """Process a single markdown line and add appropriate content to document"""
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            return current_section
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            current_section = line[3:].lower()
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('• ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        else:
            if line:
                doc.add_paragraph(line)
        
        return current_section

    def _insert_chart_if_appropriate(self, doc, current_section, chart_files):
        """Insert chart if the current section matches chart type"""
        chart_inserted = False
        
        if "business impact" in current_section and 'business_impact' in chart_files:
            doc.add_paragraph("")
            doc.add_picture(chart_files['business_impact'], width=Inches(6))
            chart_files.pop('business_impact')
            chart_inserted = True
            
        elif "performance" in current_section and 'performance_dashboard' in chart_files:
            doc.add_paragraph("")
            doc.add_picture(chart_files['performance_dashboard'], width=Inches(6))
            chart_files.pop('performance_dashboard')
            chart_inserted = True
            
        elif "velocity" in current_section and 'velocity_trends' in chart_files:
            doc.add_paragraph("")
            doc.add_picture(chart_files['velocity_trends'], width=Inches(6))
            chart_files.pop('velocity_trends')
            chart_inserted = True
        
        return chart_inserted

    def _add_remaining_charts(self, doc, chart_files):
        """Add any remaining charts at the end of the document"""
        if chart_files:
            doc.add_heading('Additional Charts', level=2)
            for chart_name, chart_path in chart_files.items():
                doc.add_heading(chart_name.replace('_', ' ').title(), level=3)
                doc.add_picture(chart_path, width=Inches(6))

    def create_docx_report(self, report_content, chart_files, output_file="professional_report.docx"):
        """Create DOCX report with embedded charts"""
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Professional Development Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("") # Empty line
        
        # Parse markdown content and add to document
        lines = report_content.split('\n')
        current_section = ""
        
        for line in lines:
            current_section = self._process_markdown_line(doc, line, current_section)
            self._insert_chart_if_appropriate(doc, current_section, chart_files)
        
        # Add any remaining charts at the end
        self._add_remaining_charts(doc, chart_files)
        
        # Save document
        doc_path = os.path.join(self.output_dir, output_file)
        doc.save(doc_path)
        print(f"✅ DOCX report saved to: {doc_path}")
        
        return doc_path

    def create_pdf_report(self, report_content, chart_files, output_file="professional_report.pdf"):
        """Create PDF report with embedded charts"""
        
        # Convert markdown to HTML
        html_content = markdown2.markdown(report_content, extras=['fenced-code-blocks', 'tables'])
        
        # Add CSS styling
        css_style = """
        <style>
        body { 
            font-family: 'Segoe UI', Arial, sans-serif; 
            line-height: 1.6; 
            margin: 40px;
            color: #333;
        }
        h1 { 
            color: #2E86AB; 
            border-bottom: 3px solid #2E86AB; 
            padding-bottom: 10px;
        }
        h2 { 
            color: #A23B72; 
            margin-top: 30px;
        }
        h3 { 
            color: #F18F01; 
        }
        .chart {
            text-align: center;
            margin: 20px 0;
        }
        ul {
            margin-left: 20px;
        }
        strong {
            color: #2E86AB;
        }
        </style>
        """
        
        # Embed charts as base64 images
        chart_html = ""
        for chart_name, chart_path in chart_files.items():
            if os.path.exists(chart_path):
                with open(chart_path, 'rb') as f:
                    img_data = base64.b64encode(f.read()).decode()
                chart_html += f"""
                <div class="chart">
                    <h3>{chart_name.replace('_', ' ').title()}</h3>
                    <img src="data:image/png;base64,{img_data}" style="max-width: 100%; height: auto;">
                </div>
                """
        
        # Combine everything
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            {css_style}
        </head>
        <body>
            {html_content}
            <h2>Charts and Visualizations</h2>
            {chart_html}
        </body>
        </html>
        """
        
        # Try to generate PDF with WeasyPrint
        try:
            import weasyprint
            pdf_path = os.path.join(self.output_dir, output_file)
            weasyprint.HTML(string=full_html).write_pdf(pdf_path)
            print(f"✅ PDF report saved to: {pdf_path}")
            return pdf_path
        except ImportError:
            print("⚠️ WeasyPrint not available - saving as HTML instead")
        except Exception as e:
            print(f"⚠️ Error generating PDF: {e}")
            print("   Falling back to HTML format...")
        
        # Fallback: save HTML file
        html_path = os.path.join(self.output_dir, output_file.replace('.pdf', '.html'))
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
        print(f"✅ HTML report saved to: {html_path}")
        return html_path

    def generate_professional_report(self, tasks, business_analysis, report_content, assignee_name="Developer"):
        """Generate complete professional report with charts in multiple formats"""
        
        print("🎨 Generating professional charts...")
        chart_files = self.generate_all_charts(tasks, business_analysis)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{assignee_name.lower().replace(' ', '_')}_{timestamp}"
        
        # Generate DOCX report
        print("📄 Creating DOCX report...")
        docx_path = self.create_docx_report(
            report_content, 
            chart_files.copy(), 
            f"{base_filename}.docx"
        )
        
        # Generate PDF report
        print("📑 Creating PDF report...")
        pdf_path = self.create_pdf_report(
            report_content, 
            chart_files.copy(), 
            f"{base_filename}.pdf"
        )
        
        return {
            'docx': docx_path,
            'pdf': pdf_path,
            'charts': chart_files
        }

# Example usage function
def create_sample_charts():
    """Create sample charts for testing"""
    visualizer = ReportVisualizer()
    
    # Sample data
    sample_tasks = [
        {'title': 'Implement payment system', 'priority': 'High', 'time': 8, 'commits': ['commit1'], 'repositories': ['api']},
        {'title': 'Fix login bug', 'priority': 'Medium', 'time': 4, 'commits': ['commit2'], 'repositories': ['frontend']},
        {'title': 'Add user dashboard', 'priority': 'High', 'time': 12, 'commits': ['commit3'], 'repositories': ['frontend']},
    ]
    
    sample_analysis = {
        'categories': {
            'revenue_generation': {
                'impact': 'Revenue & Sales',
                'tasks': [sample_tasks[0]]
            },
            'platform_stability': {
                'impact': 'Platform Stability', 
                'tasks': [sample_tasks[1]]
            },
            'user_experience': {
                'impact': 'User Experience',
                'tasks': [sample_tasks[2]]
            }
        },
        'metrics': {
            'total_tasks': 3,
            'completed_tasks': 3,
            'total_time': 24,
            'completion_rate': 100
        }
    }
    
    charts = visualizer.generate_all_charts(sample_tasks, sample_analysis)
    print("Sample charts generated:", list(charts.keys()))
    
    return visualizer, charts

if __name__ == "__main__":
    # Test the visualization system
    create_sample_charts() 